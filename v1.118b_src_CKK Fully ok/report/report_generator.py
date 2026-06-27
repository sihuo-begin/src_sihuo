# ──────────────────────────────────────────────
#  Report Generator
#  Produces a Word (.docx) report with tables and charts
# ──────────────────────────────────────────────
import logging
import os
from pathlib import Path
from datetime import datetime
from typing import Dict, List

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from utils.config import APP_NAME, APP_VERSION, CPK_SPECS, OUTPUT_DIR

logger = logging.getLogger(__name__)


class ReportGenerator:
    """Generates GRR + CPK Word report."""

    def __init__(self, output_dir: Path = OUTPUT_DIR):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def generate(
        self,
        results: Dict,        # {item: {"grr": GRRResult, "cpk": CPKResult}}
        df,                   # raw dataframe (for basic stats)
        minitab_path: str = None,
        inline_charts: bool = True,
    ) -> str:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename  = f"M600099_GRR_CPK_Report_{timestamp}.docx"
        out_path  = self.output_dir / filename

        doc = Document()

        # ── Document styles ──
        self._set_page_margins(doc)
        self._add_title(doc)
        self._add_meta(doc, results)

        # ── GRR section ──
        grr_items = {k: v for k, v in results.items() if v.get("grr")}
        if grr_items:
            self._add_grr_summary(doc, grr_items)
            self._add_grr_details(doc, grr_items, inline_charts)

        # ── CPK section ──
        cpk_items = {k: v for k, v in results.items() if v.get("cpk")}
        if cpk_items:
            self._add_cpk_summary(doc, cpk_items)
            self._add_cpk_details(doc, cpk_items, inline_charts)

        # ── Save ──
        doc.save(str(out_path))
        logger.info(f"Report saved to {out_path}")
        return str(out_path)

    # ── Helpers ─────────────────────────────────

    def _set_page_margins(self, doc):
        for section in doc.sections:
            section.top_margin    = Cm(2.0)
            section.bottom_margin = Cm(2.0)
            section.left_margin   = Cm(2.5)
            section.right_margin  = Cm(2.5)

    def _add_title(self, doc):
        p = doc.add_heading("M600099 GRR & CPK Analysis Report", level=0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.color.rgb = RGBColor(26, 42, 108)

        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sub.add_run(f"{APP_NAME}  v{APP_VERSION}  |  Generated {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(100, 100, 100)
        doc.add_paragraph()

    def _add_meta(self, doc, results: dict):
        items = list(results.keys())
        p = doc.add_paragraph()
        p.add_run("Analysis Summary\n").bold = True
        p.add_run(f"  Items analyzed:   {len(items)}\n")
        p.add_run(f"  GRR items:        {sum(1 for r in results.values() if r.get('grr'))}\n")
        p.add_run(f"  CPK items:        {sum(1 for r in results.values() if r.get('cpk'))}\n")
        p.add_run(f"  LED items:        {', '.join(items)}\n")
        doc.add_paragraph()

    def _cell_style(self, cell, text="", bold=False, bg=None, color=None, size=9):
        cell.text = text
        para = cell.paragraphs[0]
        run  = para.runs[0] if para.runs else para.add_run(text)
        run.font.size  = Pt(size)
        run.font.bold  = bold
        if color:
            run.font.color.rgb = RGBColor(*color)
        if bg:
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd  = OxmlElement("w:shd")
            shd.set(qn("w:val"),   "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"),  bg)
            tcPr.append(shd)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    def _add_table_header(self, table, headers: List[str], bg="1a2a6c"):
        row = table.rows[0]
        for i, h in enumerate(headers):
            self._cell_style(row.cells[i], h, bold=True, bg=bg, color=(255,255,255))

    # ── GRR Section ─────────────────────────────

    def _add_grr_summary(self, doc, grr_items: dict):
        doc.add_heading("1. GRR Analysis Summary", level=1)
        doc.add_paragraph(
            "GRR (Gauge Repeatability & Reproducibility) evaluates measurement system variation. "
            "GRR% = Total Variation / Tolerance.  < 10% = Excellent, 10–30% = Acceptable, > 30% = Marginal."
        )

        headers = ["Test Item", "GRR%", "P/T Ratio", "NDC", "Grade"]
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        self._add_table_header(table, headers)

        grade_colors = {
            "Excellent": "c8e6c9", "Acceptable": "fff9c4",
            "Marginal": "ffcdd2", "N/A": "eceff1"
        }

        for item, res in grr_items.items():
            grr = res.get("grr", {})
            row = table.add_row()
            cells = row.cells
            grade = grr.get("grade", "N/A")
            grade_bg = next(
                (v for k, v in grade_colors.items() if k.lower() in grade.lower()),
                "eceff1"
            )
            self._cell_style(cells[0], item)
            self._cell_style(cells[1], f"{grr.get('grr_pct', 'N/A')}%")
            self._cell_style(cells[2], str(grr.get("pt_ratio", "N/A")))
            self._cell_style(cells[3], str(grr.get("ndc", "N/A")))
            self._cell_style(cells[4], grade, bold=True, bg=grade_bg)

        doc.add_paragraph()

    def _add_grr_details(self, doc, grr_items: dict, inline: bool):
        doc.add_heading("2. GRR Detail Results", level=1)

        for item, res in grr_items.items():
            grr = res.get("grr", {})
            doc.add_heading(f"  {item}", level=2)

            table = doc.add_table(rows=7, cols=2)
            table.style = "Table Grid"
            rows_data = [
                ("Equipment Variation (EV)",     f"{grr.get('ev', 'N/A'):.4f}" if grr.get("ev") else "N/A"),
                ("Part Variation (PV)",          f"{grr.get('pv', 'N/A'):.4f}" if grr.get("pv") else "N/A"),
                ("Total Variation (TV)",         f"{grr.get('tv', 'N/A'):.4f}" if grr.get("tv") else "N/A"),
                ("Tolerance",                    f"{grr.get('tolerance', 'N/A'):.2f}" if grr.get("tolerance") else "N/A"),
                ("GRR%",                        f"{grr.get('grr_pct', 'N/A'):.2f}%" if grr.get("grr_pct") else "N/A"),
                ("P/T Ratio",                   str(grr.get("pt_ratio", "N/A"))),
                ("NDC",                         str(grr.get("ndc", "N/A"))),
            ]
            for i, (k, v) in enumerate(rows_data):
                self._cell_style(table.rows[i].cells[0], k, bold=True, bg="e8eaf6")
                self._cell_style(table.rows[i].cells[1], v)

            # Insert chart
            charts = grr.get("chart_paths", {})
            if inline and charts.get("grr_study") and os.path.isfile(charts["grr_study"]):
                try:
                    doc.add_picture(charts["grr_study"], width=Inches(5.5))
                    doc.add_paragraph("(GRR Study Chart – Minitab)")
                except Exception as e:
                    logger.warning(f"Could not insert GRR chart for {item}: {e}")

            doc.add_paragraph()

    # ── CPK Section ─────────────────────────────

    def _add_cpk_summary(self, doc, cpk_items: dict):
        doc.add_heading("3. CPK Analysis Summary", level=1)
        doc.add_paragraph(
            "CPK (Process Capability Index) measures how well the process output stays within spec limits.  "
            "CPK ≥ 1.33 = Minimum acceptable, ≥ 1.67 = Excellent, ≥ 2.0 = World-class."
        )

        headers = ["Test Item", "LSL", "USL", "Mean", "σ(within)", "Cpk", "Ppk", "Out_of_spec%", "Grade"]
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        self._add_table_header(table, headers)

        grade_colors = {
            "World-Class": "c8e6c9", "Excellent": "dcedc8",
            "Acceptable": "fff9c4", "Poor": "ffcdd2", "N/A": "eceff1"
        }

        for item, res in cpk_items.items():
            cpk = res.get("cpk", {})
            row = table.add_row()
            cells = row.cells
            grade = cpk.get("grade", "N/A")
            grade_bg = next(
                (v for k, v in grade_colors.items() if k.lower() in grade.lower()),
                "eceff1"
            )
            vals = [
                item,
                str(cpk.get("lsl", "N/A")),
                str(cpk.get("usl", "N/A")),
                f"{cpk.get('mean', 'N/A'):.2f}" if cpk.get("mean") else "N/A",
                f"{cpk.get('std_w', 'N/A'):.2f}" if cpk.get("std_w") else "N/A",
                f"{cpk.get('cpk', 'N/A'):.4f}" if cpk.get("cpk") else "N/A",
                f"{cpk.get('ppk', 'N/A'):.4f}" if cpk.get("ppk") else "N/A",
                f"{cpk.get('pct_total', 'N/A'):.4f}%" if cpk.get("pct_total") else "N/A",
                grade,
            ]
            for i, v in enumerate(vals):
                bold = (i == len(vals) - 1)
                bg   = grade_bg if bold else None
                self._cell_style(cells[i], v, bold=bold, bg=bg)

        doc.add_paragraph()

    def _add_cpk_details(self, doc, cpk_items: dict, inline: bool):
        doc.add_heading("4. CPK Detail Results", level=1)

        for item, res in cpk_items.items():
            cpk = res.get("cpk", {})
            doc.add_heading(f"  {item}", level=2)

            table = doc.add_table(rows=12, cols=2)
            table.style = "Table Grid"
            rows_data = [
                ("LSL",            str(cpk.get("lsl", "N/A"))),
                ("USL",            str(cpk.get("usl", "N/A"))),
                ("Mean",           f"{cpk.get('mean', 'N/A'):.4f}" if cpk.get("mean") else "N/A"),
                ("σ (within)",     f"{cpk.get('std_w', 'N/A'):.4f}" if cpk.get("std_w") else "N/A"),
                ("σ (overall)",    f"{cpk.get('std',  'N/A'):.4f}" if cpk.get("std")  else "N/A"),
                ("Cp",             f"{cpk.get('cp',   'N/A'):.4f}" if cpk.get("cp")   else "N/A"),
                ("Cpl",            f"{cpk.get('cpl',  'N/A'):.4f}" if cpk.get("cpl")  else "N/A"),
                ("Cpu",            f"{cpk.get('cpu',  'N/A'):.4f}" if cpk.get("cpu")  else "N/A"),
                ("Cpk",            f"{cpk.get('cpk',  'N/A'):.4f}" if cpk.get("cpk")  else "N/A"),
                ("Ppk",            f"{cpk.get('ppk',  'N/A'):.4f}" if cpk.get("ppk")  else "N/A"),
                ("% below LSL",   f"{cpk.get('pct_outside_lsl', 'N/A'):.4f}%" if cpk.get("pct_outside_lsl") is not None else "N/A"),
                ("% above USL",   f"{cpk.get('pct_outside_usl', 'N/A'):.4f}%" if cpk.get("pct_outside_usl") is not None else "N/A"),
            ]
            for i, (k, v) in enumerate(rows_data):
                self._cell_style(table.rows[i].cells[0], k, bold=True, bg="e8eaf6")
                self._cell_style(table.rows[i].cells[1], v)

            charts = cpk.get("chart_paths", {})
            if inline and charts.get("capability") and os.path.isfile(charts["capability"]):
                try:
                    doc.add_picture(charts["capability"], width=Inches(5.5))
                    doc.add_paragraph("(Capability Histogram – Minitab)")
                except Exception as e:
                    logger.warning(f"Could not insert CPK chart for {item}: {e}")

            doc.add_paragraph()
